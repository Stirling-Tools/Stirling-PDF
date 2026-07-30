"""Fill DOCX templates from JSON data: ``{{ dotted.path }}`` placeholders.

Scalar placeholders are replaced everywhere (body, tables, headers, footers).
A table row whose text contains ``{{#items.field}}`` markers is treated as a
row template: it is cloned once per element of the ``items`` array. Unresolved
placeholders are left in place and reported back so the caller can surface them.

Formatting caveat: a placeholder split across styled runs collapses that
paragraph's text into its first run's style.
"""

from __future__ import annotations

import base64
import copy
import io
import re
from typing import Any

from pydantic import JsonValue

from stirling.contracts.docparse import FillDocxRequest, FillDocxResponse

_PLACEHOLDER = re.compile(r"\{\{\s*(#?[\w.]+)\s*\}\}")


def _resolve(path: str, data: dict[str, Any]) -> Any | None:
    node: Any = data
    for part in path.split("."):
        if isinstance(node, dict) and part in node:
            node = node[part]
        else:
            return None
    return node


def _render_scalar(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, list):
        return ", ".join(_render_scalar(v) for v in value)
    return str(value)


class _Stats:
    def __init__(self) -> None:
        self.replaced = 0
        self.missing: set[str] = set()


def _fill_paragraph(paragraph: Any, data: dict[str, Any], stats: _Stats) -> None:
    text = paragraph.text
    if "{{" not in text:
        return

    def substitute(match: re.Match[str]) -> str:
        path = match.group(1)
        if path.startswith("#"):
            return match.group(0)  # row-template marker, handled at table level
        value = _resolve(path, data)
        if value is None:
            stats.missing.add(path)
            return match.group(0)
        stats.replaced += 1
        return _render_scalar(value)

    rendered = _PLACEHOLDER.sub(substitute, text)
    if rendered == text:
        return
    # Collapse into the first run to survive placeholders split across runs.
    if paragraph.runs:
        paragraph.runs[0].text = rendered
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(rendered)


def _row_template_array(row: Any) -> str | None:
    """Return the array name when the row carries ``{{#name.field}}`` markers."""
    names = {
        match.group(1)[1:].split(".")[0]
        for cell in row.cells
        for match in _PLACEHOLDER.finditer(cell.text)
        if match.group(1).startswith("#")
    }
    return names.pop() if len(names) == 1 else None


def _fill_table(table: Any, data: dict[str, Any], stats: _Stats) -> None:
    for row in list(table.rows):
        array_name = _row_template_array(row)
        if array_name is None:
            continue
        items = _resolve(array_name, data)
        if not isinstance(items, list):
            stats.missing.add(array_name)
            continue
        for _ in items:
            new_row = copy.deepcopy(row._tr)
            row._tr.addprevious(new_row)
        # Clones sit before the template; rewrite their markers, then drop the template.
        _rewrite_cloned_rows(table, row, array_name, items, data, stats)
        row._tr.getparent().remove(row._tr)


def _rewrite_cloned_rows(
    table: Any, template_row: Any, array_name: str, items: list[Any], data: dict[str, Any], stats: _Stats
) -> None:
    marker_prefix = f"#{array_name}"
    clones = [
        r for r in table.rows if r._tr is not template_row._tr and marker_prefix in "".join(c.text for c in r.cells)
    ]
    for row, item in zip(clones, items, strict=False):
        scoped = dict(data)
        scoped[array_name] = item if isinstance(item, dict) else {"value": item}
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text

                def substitute(match: re.Match[str]) -> str:
                    path = match.group(1)
                    if not path.startswith(marker_prefix):
                        return match.group(0)
                    item_path = path[1:]  # "#items.field" -> "items.field"
                    value = _resolve(item_path, scoped)
                    if value is None and "." not in item_path:
                        value = scoped[array_name].get("value") if isinstance(scoped[array_name], dict) else None
                    if value is None:
                        stats.missing.add(item_path)
                        return match.group(0)
                    stats.replaced += 1
                    return _render_scalar(value)

                rendered = _PLACEHOLDER.sub(substitute, text)
                if rendered != text:
                    if paragraph.runs:
                        paragraph.runs[0].text = rendered
                        for run in paragraph.runs[1:]:
                            run.text = ""
                    else:
                        paragraph.add_run(rendered)


def _walk_paragraphs(document: Any) -> list[tuple[Any, Any]]:
    """Yield (paragraph, containing table or None) across body, tables, headers, footers."""
    found: list[tuple[Any, Any]] = [(p, None) for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                found.extend((p, table) for p in cell.paragraphs)
    for section in document.sections:
        for part in (section.header, section.footer):
            found.extend((p, None) for p in part.paragraphs)
    return found


def fill_docx(request: FillDocxRequest) -> FillDocxResponse:
    import docx  # local import: python-docx is small but only needed here

    data: dict[str, JsonValue] = dict(request.data)
    document = docx.Document(io.BytesIO(base64.b64decode(request.template_base64)))
    stats = _Stats()

    for table in document.tables:
        _fill_table(table, data, stats)
    for paragraph, _table in _walk_paragraphs(document):
        _fill_paragraph(paragraph, data, stats)

    out = io.BytesIO()
    document.save(out)
    return FillDocxResponse(
        docx_base64=base64.b64encode(out.getvalue()).decode("ascii"),
        replaced=stats.replaced,
        missing=sorted(stats.missing),
    )
