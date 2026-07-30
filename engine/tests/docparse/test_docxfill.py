from __future__ import annotations

import base64
import io
from typing import Any

import docx

from stirling.contracts.docparse import FillDocxRequest
from stirling.docparse.docxfill import fill_docx


def _template_base64() -> str:
    document = docx.Document()
    document.add_paragraph("Dear {{ customer.name }},")
    document.add_paragraph("Your total is {{ total }}.")
    document.add_paragraph("Unknown: {{ nowhere.field }}")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Price"
    table.rows[1].cells[0].text = "{{#items.name}}"
    table.rows[1].cells[1].text = "{{#items.price}}"
    buffer = io.BytesIO()
    document.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def _load(response_base64: str) -> Any:
    return docx.Document(io.BytesIO(base64.b64decode(response_base64)))


def test_fills_scalars_tables_and_reports_missing() -> None:
    request = FillDocxRequest(
        template_base64=_template_base64(),
        data={
            "customer": {"name": "ACME GmbH"},
            "total": 12.5,
            "items": [
                {"name": "Widget", "price": "2.00"},
                {"name": "Gadget", "price": "10.50"},
            ],
        },
    )
    response = fill_docx(request)
    filled = _load(response.docx_base64)

    paragraphs = [p.text for p in filled.paragraphs]
    assert "Dear ACME GmbH," in paragraphs
    assert "Your total is 12.5." in paragraphs
    # Unresolved placeholders stay put and are reported.
    assert any("{{ nowhere.field }}" in p for p in paragraphs)
    assert response.missing == ["nowhere.field"]

    table = filled.tables[0]
    rendered_rows = [[cell.text for cell in row.cells] for row in table.rows]
    assert ["Widget", "2.00"] in rendered_rows
    assert ["Gadget", "10.50"] in rendered_rows
    # The template row is gone.
    assert all("{{#" not in cell for row in rendered_rows for cell in row)
    assert response.replaced >= 6


def test_empty_items_removes_template_row() -> None:
    request = FillDocxRequest(
        template_base64=_template_base64(),
        data={"customer": {"name": "X"}, "total": 1, "items": []},
    )
    response = fill_docx(request)
    filled = _load(response.docx_base64)
    assert len(filled.tables[0].rows) == 1  # only the header remains
